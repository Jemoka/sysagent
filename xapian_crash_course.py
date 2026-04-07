"""
Xapian Crash Course for Python
==============================

Install: sudo apt install python3-xapian
  (pip doesn't work well — xapian bindings are C++ and need the
   system lib. On mac: brew install xapian --with-python3)

Core concepts:
  - Database: a directory on disk holding the index
  - Document: a thing you index (has terms, values, and data)
  - Terms: searchable tokens, stored with POSITIONS for proximity search
  - TermGenerator: tokenizes text into terms with positions automatically
  - QueryParser: parses human-readable queries into Query objects
  - Enquire: runs queries against a database
"""

import xapian
import os
import json
import hashlib


# =============================================================================
# PART 1: INDEXING
# =============================================================================

def create_or_open_db(db_path: str) -> xapian.WritableDatabase:
    """Open (or create) a writable database. That's it."""
    return xapian.WritableDatabase(db_path, xapian.DB_CREATE_OR_OPEN)


def index_document(db: xapian.WritableDatabase, doc_id: str, text: str,
                   metadata: dict | None = None):
    """
    Index a single document idempotently.

    doc_id:   unique identifier (e.g. file path). Used as the idempotency key.
    text:     the extracted full text to index.
    metadata: optional dict stored as JSON in the document data.
    """
    # We use a "unique ID term" prefixed with Q to enable replace-or-insert.
    # This is the standard Xapian pattern for idempotent upserts.
    id_term = "Q" + doc_id

    doc = xapian.Document()

    # TermGenerator handles tokenization, stemming, and — critically —
    # storing term POSITIONS, which is what enables phrase/proximity search.
    termgen = xapian.TermGenerator()
    termgen.set_stemmer(xapian.Stem("en"))
    termgen.set_document(doc)

    # index_text() adds terms WITH positions. This is the important part.
    # Without positions, you get boolean search only — no NEAR, no phrases.
    termgen.index_text(text)

    # You can index specific fields with prefixes for field-scoped search.
    # Convention: prefixes are uppercase letters.
    if metadata:
        if "title" in metadata:
            termgen.index_text(metadata["title"], 1, "S")  # S = subject/title
        if "author" in metadata:
            termgen.index_text(metadata["author"], 1, "A")

    # Store the unique ID term for idempotent replace
    doc.add_boolean_term(id_term)

    # Store whatever you want to retrieve later as document "data".
    # This is NOT searched — it's just payload.
    doc.set_data(json.dumps({
        "id": doc_id,
        "metadata": metadata or {},
    }))

    # replace_document with the id_term: inserts if new, replaces if exists.
    # This is your idempotency.
    db.replace_document(id_term, doc)


def index_folder(db_path: str, folder: str, extractors: dict | None = None):
    """
    Walk a folder, extract text, index everything idempotently.

    extractors: dict mapping file extension -> callable(filepath) -> str
                If not provided, only .txt files are indexed.
    """
    if extractors is None:
        extractors = {".txt": lambda p: open(p).read()}

    db = create_or_open_db(db_path)

    count = 0
    for root, dirs, files in os.walk(folder):
        for fname in files:
            ext = os.path.splitext(fname)[1].lower()
            if ext not in extractors:
                continue

            filepath = os.path.join(root, fname)
            stat = os.stat(filepath)

            # Use path + mtime + size as a change-detection key.
            # If nothing changed, we still call replace_document (it's cheap)
            # but you could skip it entirely for speed.
            doc_id = filepath

            try:
                text = extractors[ext](filepath)
            except Exception as e:
                print(f"SKIP {filepath}: {e}")
                continue

            index_document(db, doc_id, text, metadata={
                "title": fname,
                "path": filepath,
                "size": stat.st_size,
                "mtime": stat.st_mtime,
            })
            count += 1

    # Commit
    db.close()
    print(f"Indexed {count} documents into {db_path}")


# =============================================================================
# PART 2: SEARCHING
# =============================================================================

def search(db_path: str, query_string: str, offset: int = 0,
           limit: int = 10) -> list[dict]:
    """
    Search the index. Returns list of {rank, score, data} dicts.

    query_string supports:
      - Simple terms:          contract liability
      - Phrases:               "breach of contract"
      - Proximity (NEAR):      contract NEAR/5 liability
      - Boolean:               contract AND liability NOT damages
      - Field-scoped:          title:contract   author:smith
      - Wildcards:             contra*
      - Combined:              "data breach" NEAR/10 notification AND title:privacy

    The NEAR operator is what you care about:
      term1 NEAR/N term2  ->  terms within N words of each other, either order
    """
    db = xapian.Database(db_path)  # read-only

    # QueryParser turns human-readable strings into Query objects.
    qp = xapian.QueryParser()
    qp.set_stemmer(xapian.Stem("en"))
    qp.set_stemming_strategy(xapian.QueryParser.STEM_SOME)
    qp.set_database(db)  # needed for wildcard expansion

    # Register field prefixes so "title:foo" and "author:bar" work
    qp.add_prefix("title", "S")
    qp.add_prefix("author", "A")

    # Enable all the good stuff
    flags = (
        xapian.QueryParser.FLAG_BOOLEAN |        # AND, OR, NOT
        xapian.QueryParser.FLAG_PHRASE |          # "exact phrase"
        xapian.QueryParser.FLAG_LOVEHATE |        # +required -excluded
        xapian.QueryParser.FLAG_WILDCARD |        # prefix*
        xapian.QueryParser.FLAG_NEAR |            # NEAR/n  <-- the money flag
        xapian.QueryParser.FLAG_PARTIAL            # partial matching
    )

    query = qp.parse_query(query_string, flags)
    print(f"Parsed query: {query}")

    # Enquire runs the query
    enquire = xapian.Enquire(db)
    enquire.set_query(query)

    results = []
    for match in enquire.get_mset(offset, limit):
        data = json.loads(match.document.get_data())
        results.append({
            "rank": match.rank + 1,
            "score": match.weight,
            "doc_id": data.get("id"),
            "metadata": data.get("metadata", {}),
        })

    return results


# =============================================================================
# PART 3: DOCUMENT EXTRACTORS (plug in what you need)
# =============================================================================

def make_extractors():
    """
    Build a dict of file extension -> text extractor functions.
    Add whatever you have installed.
    """
    extractors = {
        ".txt": lambda p: open(p, errors="replace").read(),
        ".md":  lambda p: open(p, errors="replace").read(),
    }

    # PDF via pdfplumber (pip install pdfplumber)
    try:
        import pdfplumber
        def extract_pdf(path):
            with pdfplumber.open(path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        extractors[".pdf"] = extract_pdf
    except ImportError:
        pass

    # DOCX via python-docx (pip install python-docx)
    try:
        import docx
        def extract_docx(path):
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        extractors[".docx"] = extract_docx
    except ImportError:
        pass

    # HTML via beautifulsoup
    try:
        from bs4 import BeautifulSoup
        def extract_html(path):
            with open(path, errors="replace") as f:
                return BeautifulSoup(f.read(), "html.parser").get_text()
        extractors[".html"] = extract_html
        extractors[".htm"] = extract_html
    except ImportError:
        pass

    return extractors


# =============================================================================
# PART 4: PUTTING IT TOGETHER
# =============================================================================

if __name__ == "__main__":
    import sys

    DB_PATH = "./my_index.db"

    if len(sys.argv) < 2:
        print("Usage:")
        print("  python xapian_crash_course.py index /path/to/folder")
        print("  python xapian_crash_course.py search 'contract NEAR/5 liability'")
        print("  python xapian_crash_course.py search '\"data breach\" AND notification'")
        sys.exit(1)

    cmd = sys.argv[1]

    if cmd == "index":
        folder = sys.argv[2]
        index_folder(DB_PATH, folder, make_extractors())

    elif cmd == "search":
        query = " ".join(sys.argv[2:])
        results = search(DB_PATH, query)
        if not results:
            print("No results.")
        for r in results:
            print(f"  #{r['rank']} (score: {r['score']:.2f}) {r['doc_id']}")
            if r["metadata"]:
                print(f"         {r['metadata']}")

    elif cmd == "delete":
        # Bonus: remove a doc by its id
        doc_id = sys.argv[2]
        db = create_or_open_db(DB_PATH)
        db.delete_document("Q" + doc_id)
        db.close()
        print(f"Deleted {doc_id}")

    else:
        print(f"Unknown command: {cmd}")
