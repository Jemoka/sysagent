from docx import Document as DocxDocument


def _extract_docx(path):
    document = DocxDocument(path)
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n".join(parts)
